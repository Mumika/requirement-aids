import json
from django.shortcuts import render, redirect
from firstWEB import models
from django.http import HttpResponse
from docx import Document
import time
import re
# 全局变量
memberpos='成员'


# 登陆验证
def login(request):
    global memberpos
    if request.method == "POST":
        logname = request.POST.get('logname')
        logpwd = request.POST.get('logpwd')
        empwd = models.members.objects.get(empno=logname).empwd
        logpos=models.members.objects.get(empno=logname).pos
        memberpos=logpos
        if empwd != logpwd:
            return HttpResponse('密码或账号错误')
        else:
            return redirect('/index/')
    return render(request, "login.html")

def index(request):
    return render(request, "index.html")


# index显示
def left(request):
    return render(request, 'left.html')
def head(request):
    return render(request, 'head.html')
def head2(request):
    return render(request, 'head2.html')
def main(request):
    return render(request, 'main.html')


# 成员，任务，日志数据库相关
def addMember(request):
    global memberpos
    if request.method == "POST" and memberpos == '经理':
        empno = request.POST.get("empno")
        empwd = request.POST.get("empwd")
        empname = request.POST.get("empname")
        age = request.POST.get("age")
        sex = request.POST.get("sex")
        pos = request.POST.get("pos")
        group = request.POST.get("group")
        models.members.objects.create(empno=empno, empwd=empwd, empname=empname, age=age, sex=sex, pos=pos, group=group)
        return HttpResponse('添加成功')
    elif memberpos != '经理':
        return HttpResponse('您没有权限')
    return render(request, "addmember.html")


def memberList(request):
    memberList = models.members.objects.all()
    global memberpos
    if request.method == "POST" and memberpos == '经理':
        for emp in memberList:
            emppos = request.POST.get('memberpos' + emp.empno)
            empgroup = request.POST.get('membergroup' + emp.empno)
            models.members.objects.filter(empno=emp.empno).update(pos=emppos)
            models.members.objects.filter(empno=emp.empno).update(group=empgroup)
            if request.POST.get('delmember' + emp.empno) == '移除':
                models.members.objects.filter(empno=emp.empno).delete()
        return HttpResponse('更新信息成功')
    elif memberpos != '经理':
        return HttpResponse('您没有权限')
    return render(request, "memberlist.html", context={"member_obj_list": memberList})



# 加备注
def addtaskremarks(request):
    if request.method == 'POST':
        taskno = request.POST.get("taskno")
        task = models.task.objects.get(taskno=taskno)
        remarks = request.POST.get("remarks")
        task.remarks = task.remarks + ' ' + remarks
        task.save()
    return render(request, 'addtaskremarks.html')


def addLog(request):
    hdList=models.hd.objects.all()
    if request.method == "POST":
        stakeholder = request.POST.get("stakeholder")
        logno = request.POST.get("logno")
        logname = request.POST.get("logname")
        doc = request.POST.get("doc")
        remarks = request.POST.get("remarks")
        # 保存到数据库
        models.log.objects.create(stakeholder=stakeholder, logno=logno, logname=logname, doc=doc, remarks=remarks)
        return HttpResponse('添加成功')
    return render(request, "addlog.html",context={"hd_obj_list":hdList})


def logList(request):
    logList = models.log.objects.all()
    if request.method == "POST":
        for log in logList:
            if request.POST.get('dellog' + log.logno) == '移除':
                models.log.objects.filter(logno=log.logno).delete()
    return render(request, "loglist.html", context={"log_obj_list": logList})


def addTask(request):
    if request.method == "POST":
        taskno = request.POST.get("taskno")
        taskname = request.POST.get("taskname")
        time = request.POST.get("time")
        ddl = request.POST.get("ddl")
        taskgroup = request.POST.get("taskgroup")
        taskperson = request.POST.get("taskperson")
        state = request.POST.get('state')
        remarks = request.POST.get("remarks")
        # 保存到数据库
        models.task.objects.create(taskno=taskno, taskname=taskname, time=time, ddl=ddl, taskgroup=taskgroup,
                                   taskperson=taskperson, state=state, remarks=remarks)
        return HttpResponse('添加成功')
    return render(request, "addtask.html")


def taskList(request):
    global memberpos
    taskList = models.task.objects.all()
    if request.method == "POST":
        if memberpos == '经理' or memberpos == '组长':
            for task in taskList:
                    taskgroup = request.POST.get('taskgroup' + task.taskno)
                    taskstate = request.POST.get('taskstate' + task.taskno)
                    taskperson = request.POST.get('taskperson' + task.taskno)
                    models.task.objects.filter(taskno=task.taskno).update(taskgroup=taskgroup)
                    models.task.objects.filter(taskno=task.taskno).update(state=taskstate)
                    models.task.objects.filter(taskno=task.taskno).update(taskperson=taskperson)
            return HttpResponse('更新信息成功')
        else:
            return HttpResponse('修改失败')
    return render(request, "tasklist.html", context={"task_obj_list": taskList})


# 需求获取模块
# 干系人识别
def addhoder(request):
    hdlstlen = len(models.hd.objects.all())
    if request.method == 'POST':
        s=request.POST.get('holderlist')
        if s!= '':
            holderlist =re.split('[，,、 ]+',s)#正则表达式[,|;*]中的任何一个出现至少一次

            for i, hoder in enumerate(holderlist):
                models.hd.objects.create(hdno=hdlstlen + i + 1, hdname=hoder)
        else:
            return HttpResponse('请输入有效数据')
    return render(request, 'addhoder.html')

def hoderlist(request):
    hdlist = models.hd.objects.all()
    if request.method == "POST":
        for hd in hdlist:
            hdstyle = request.POST.get('hdstyle' + hd.hdno)
            models.hd.objects.filter(hdno=hd.hdno).update(hdstyle=hdstyle)
        return HttpResponse('添加成功')
    return render(request, "hoderlist.html", context={"hd_obj_list": hdlist})


# 选择三种需求获取方式
def selectmodel(request):
    return render(request, 'selectmodel.html')


def fieldwork(request):
    gbList = models.gb.objects.all()
    return render(request, 'fieldwork.html', context={"gb_obj_list": gbList})


def interview(request):
    if request.method == "POST":
        itviewno = request.POST.get('itviewno')
        itviewque = request.POST.get('itviewque')
        itviewans = request.POST.get('itviewans')
        models.interview.objects.create(itviewno=itviewno, itviewque=itviewque, itviewans=itviewans)
        return HttpResponse('添加成功')
    return render(request, 'interview.html')


def questionnaire(request):
    quelist1 = models.que.objects.filter(quemodel='受众范围').values_list('quename')
    quelist2 = models.que.objects.filter(quemodel='软件风格').values_list('quename')
    quelist3 = models.que.objects.filter(quemodel='性能要求').values_list('quename')
    quelist4 = models.que.objects.filter(quemodel='用户体验').values_list('quename')
    print(quelist1, quelist2, quelist3, quelist4)
    lst1, lst2, lst3, lst4 = [], [], [], []
    for i in quelist1:
        lst1.append(i)
    for i in quelist2:
        lst2.append(i)
    for i in quelist3:
        lst3.append(i)
    for i in quelist4:
        lst4.append(i)
    print(lst1, lst2, lst3, lst4)
    print(str(lst1), str(lst2), str(lst3), str(lst4))

    if request.method == 'POST':
        quename = request.POST.get('ys')
        print(quename)
        queans = request.POST.get('queans')
        models.que.objects.filter(quename=quename).update(queans=queans)
    return render(request, 'questionnaire.html',
                  context={"que_obj_list1": json.dumps(lst1), "que_obj_list2": json.dumps(lst2),
                           "que_obj_list3": json.dumps(lst3), "que_obj_list4": json.dumps(lst4)})


def addque(request):
    if request.POST.get('submitque'):
        queno = request.POST.get('queno')
        quename = request.POST.get('quename')
        quemodel = request.POST.get('quemodel')
        models.que.objects.create(queno=queno, quename=quename, quemodel=quemodel)
    return render(request, 'addque.html')


def addgb(request):
    if request.method == 'POST':
        gbno = request.POST.get('gbno')
        gbname = request.POST.get('gbname')
        models.gb.objects.create(gbno=gbno, gbname=gbname)
        return HttpResponse('添加成功')
    return render(request, 'addgb.html')


# 添加用户故事
def addStory(request):
    itviewList = models.interview.objects.all()
    gbList = models.gb.objects.all()
    queList = models.que.objects.all()
    storyList = models.story.objects.all()
    hdList=models.hd.objects.all()
    if request.method == 'POST':
        storyno = len(storyList) + 1
        stakeholder = request.POST.get('stakeholder')
        role = request.POST.get('role')
        activity = request.POST.get('activity')
        value = request.POST.get('value')
        itview = request.POST.get('interview')
        quenaire = request.POST.get('que')
        remarks = request.POST.get('remarks')
        models.story.objects.create(storyno=storyno, stakeholder=stakeholder,
                                    role=role, activity=activity, value=value,
                                    itview=itview, quenaire=quenaire, remarks=remarks)
        return HttpResponse('添加成功')
    return render(request, 'addStory.html',
                  context={"itview_obj_list": itviewList, "gb_obj_list": gbList, "que_obj_list": queList,"hd_obj_list":hdList})

def addsetting(request):
    if request.method=='POST':
        bg=request.POST.get('bg')
        goal=request.POST.get('goal')
        success=request.POST.get('success')
        remarks=request.POST.get('remarks')
        print(bg,goal,success,remarks)
        models.setting.objects.create(bg=bg,goal=goal,success=success,remarks=remarks)
        return HttpResponse('添加成功')
    else:
        return render(request,'addsetting.html')

# 生成需求报告
def requDemo(request):
    storylist = models.story.objects.all()
    funquelist=models.que.objects.filter(quemodel='性能要求')
    gblist=models.gb.objects.all()
    setlist=models.setting.objects.all()
    localtime = time.asctime(time.localtime(time.time()))
    print(storylist)
    if request.method == "POST":
        opstyle = request.POST.get('output')
        print(opstyle)
        if opstyle == 'docx':
            # 导出到doc
            document = Document()  # 打开一个基于默认“模板”的空白文档
            document.add_heading('需求规格说明书（初稿）', 0)  # 添加标题
            document.add_paragraph('导出时间:' + localtime)
            document.add_paragraph('1.前景')
            document.add_paragraph('1.1背景')
            for bg in setlist:
                document.add_paragraph("　　"+str(bg.bg))
            document.add_paragraph('1.2业务目标')
            for goal in setlist:
                document.add_paragraph("　　"+str(goal.goal))
            document.add_paragraph('1.3成功标准')
            for suc in setlist:
                document.add_paragraph("　　"+str(suc.success))
            document.add_paragraph('2.用例')
            for i in storylist:
                document.add_paragraph(i.storyno)
                document.add_paragraph('干系人：' + str(i.stakeholder))
                document.add_paragraph('角色：' + str(i.role))
                document.add_paragraph('活动：' + str(i.activity))
                document.add_paragraph('价值：' + str(i.value))
                document.add_paragraph('关联的访谈内容：' + str(i.itview))
                document.add_paragraph('关联的问卷问题：' + str(i.quenaire))
                document.add_paragraph('备注' + str(i.remarks))
            document.add_paragraph('3.非功能性需求')
            for que in funquelist:
                document.add_paragraph("　　"+str(que.quename)+str(que.queans))
            document.add_paragraph('4.备注')
            for rm in setlist:
                document.add_paragraph("　　"+str(rm.remarks))
            document.add_paragraph('5.参考文献或标准')
            for num, gb in enumerate(gblist):
                document.add_paragraph("　　"+str(num+1)+'.'+str(gb.gbname))
            document.save(r'C:\Users\ECHO\Desktop\output\output.docx')
            return HttpResponse('导出成功')
        elif opstyle == 'txt':
            # 导出到txt
            f = open(r'C:\Users\ECHO\Desktop\output\output.txt', 'w')
            f.write('需求规格说明书（初稿）')  # 添加标题
            f.write('\n导出时间:' + localtime)
            f.write('\n1.前景')
            f.write('\n1.1背景')
            for bg in setlist:
                f.write("\n　　" + str(bg.bg))
            f.write('\n1.2业务目标')
            for goal in setlist:
                f.write("\n　　" + str(goal.goal))
            f.write('\n1.3成功标准')
            for suc in setlist:
                f.write("　　" + str(suc.success))
            f.write('\n2.用户故事')
            for i in storylist:
                f.write('\n' + i.storyno)
                f.write('\n干系人：' + str(i.stakeholder))
                f.write('\n角色：' + str(i.role))
                f.write('\n活动：' + str(i.activity))
                f.write('\n价值：' + str(i.value))
                f.write('\n关联的访谈内容：' + str(i.itview))
                f.write('\n关联的问卷问题：' + str(i.quenaire))
                f.write('\n备注' + str(i.remarks))
            f.write('\n3.非功能性需求')
            for que in funquelist:
                f.write("\n　　" + str(que.quename) + str(que.queans))
            f.write('\n4.备注')
            for rm in setlist:
                f.write("\n　　" + str(rm.remarks))
            f.write('\n5.参考文献或标准')
            for gb in gblist:
                f.write("\n　　" + str(gb.gbno) + '.' + str(gb.gbname))
            f.close()
            return HttpResponse('导出成功')
    return render(request, 'requdemo.html')
